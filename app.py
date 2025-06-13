import os
import re
from datetime import datetime
from collections import Counter

from flask import (
    Flask, render_template, request,
    redirect, url_for, flash, send_from_directory
)
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, UserMixin, login_user, login_required,
    logout_user, current_user
)
from passlib.hash import pbkdf2_sha256
from flask_migrate import Migrate
from PyPDF2 import PdfReader
import docx
from docx import Document
import stripe
import openai

# --- CONFIG ---
openai.api_key = os.getenv("OPENAI_API_KEY")
stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "")  # Set in environment

app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'changeme-in-prod')
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
DB_PATH = os.path.join(BASE_DIR, 'database.db')

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{DB_PATH}"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Ensure uploads folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

db = SQLAlchemy(app)
migrate = Migrate(app, db)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

# --- MODELS ---
class User(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(256), nullable=False)
    role = db.Column(db.String(20), nullable=False)  # contractor, agency, manager
    area = db.Column(db.String(100), nullable=True)
    credits = db.Column(db.Integer, default=0)
    email = db.Column(db.String(120), nullable=True)
    address = db.Column(db.String(200), nullable=True)
    phone = db.Column(db.String(40), nullable=True)
    resumes = db.relationship('Resume', backref='user', lazy=True)

class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(260), nullable=False)
    upload_time = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    summary = db.Column(db.Text, nullable=True)
    generated_cover_letter = db.Column(db.Text, nullable=True)
    tags = db.relationship('ResumeTag', backref='resume', lazy=True, cascade="all, delete-orphan")

class ResumeTag(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    tag = db.Column(db.String(100), nullable=False)
    resume_id = db.Column(db.Integer, db.ForeignKey('resume.id'), nullable=False)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

ALLOWED_EXTENSIONS = {'pdf', 'docx'}
STOPWORDS = {'the','and','for','with','that','this','from','your','have','will','project','manager','experience'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(filepath):
    ext = filepath.rsplit('.', 1)[1].lower()
    text = ''
    if ext == 'pdf':
        reader = PdfReader(filepath)
        for page in reader.pages:
            text += page.extract_text() or ''
    elif ext == 'docx':
        docf = docx.Document(filepath)
        for p in docf.paragraphs:
            text += p.text + ' '
    return text

def parse_and_save_tags(resume):
    path = os.path.join(app.config['UPLOAD_FOLDER'], resume.filename)
    raw = extract_text(path).lower()
    words = re.findall(r'\b[a-z]{4,}\b', raw)
    words = [w for w in words if w not in STOPWORDS]
    top = [w for w, _ in Counter(words).most_common(15)]
    for w in top:
        tag = ResumeTag(tag=w, resume_id=resume.id)
        db.session.add(tag)
    db.session.commit()

def generate_cv_summary(cv_text):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an expert CV summariser."},
            {"role": "user", "content": f"Summarise this CV for recruiters: {cv_text[:3500]}"}
        ],
        max_tokens=120
    )
    return response['choices'][0]['message']['content'].strip()

# --- NEW COVER LETTER HELPER ---
def generate_cover_letter_text(cv_text: str, job_desc: str, tone: str) -> str:
    prompt = (
        f"You are an expert career coach. "
        f"Write a {tone} cover letter for a candidate whose CV is below, applying for this job:\n\n"
        f"=== CV ===\n{cv_text[:3000]}\n\n"
        f"=== Job Description ===\n{job_desc[:2000]}\n\n"
        "The letter should be 3–4 paragraphs, highlight key skills, match tone, and end with a strong closing."
    )
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=500
    )
    return response.choices[0].message.content.strip()

# --- ROUTES ---
@app.route('/')
@login_required
def home():
    if current_user.role == "contractor":
        resumes = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.upload_time.desc()).all()
        return render_template('contractor_home.html', resumes=resumes)
    elif current_user.role == "manager":
        return redirect(url_for('manager_all_cvs'))
    elif current_user.role == "agency":
        return redirect(url_for('search'))
    else:
        return render_template('index.html')

# [existing auth/register/login/logout/profile/upload/download/... routes remain unchanged]

# --- COVER LETTER GENERATION ROUTE ---
@app.route('/contractor/cover_letter/<int:resume_id>', methods=['GET', 'POST'])
@login_required
def generate_cover_letter(resume_id):
    resume = Resume.query.get_or_404(resume_id)
    if resume.user_id != current_user.id:
        flash("Not allowed.", "danger")
        return redirect(url_for('home'))

    cover_letter = None
    job_desc = ""
    if request.method == 'POST':
        job_desc = request.form['job_description'].strip()
        tone = request.form.get('tone', 'professional')

        # Credit check (1 credit per letter)
        if current_user.credits < 1:
            flash("Not enough credits. Please buy more.", "warning")
            return redirect(url_for('buy_credits'))

        # Deduct credit
        current_user.credits -= 1
        db.session.commit()

        # Load CV and generate
        path = os.path.join(app.config['UPLOAD_FOLDER'], resume.filename)
        cv_text = extract_text(path)
        try:
            cover_letter = generate_cover_letter_text(cv_text, job_desc, tone)
            resume.generated_cover_letter = cover_letter
            db.session.commit()
            flash("Cover letter generated!", "success")
        except Exception as e:
            flash(f"Cover letter generation failed: {e}", "danger")

    return render_template(
        'cover_letter.html',
        resume=resume,
        cover_letter=cover_letter,
        job_description=job_desc
    )

# --- DOWNLOAD GENERATED COVER LETTER ---
@app.route('/download_cover_letter/<int:resume_id>')
@login_required
def download_cover_letter(resume_id):
    resume = Resume.query.get_or_404(resume_id)
    if resume.user_id != current_user.id:
        flash("Not allowed.", "danger")
        return redirect(url_for('home'))

    # Convert to .docx on-the-fly
    doc = Document()
    doc.add_heading('Cover Letter', level=1)
    for line in (resume.generated_cover_letter or '').split('\n'):
        doc.add_paragraph(line)
    tmp_filename = f"cover_letter_{resume.id}.docx"
    tmp_path = os.path.join(app.config['UPLOAD_FOLDER'], tmp_filename)
    doc.save(tmp_path)
    return send_from_directory(
        app.config['UPLOAD_FOLDER'],
        tmp_filename,
        as_attachment=True
    )

# --- AGENCY & MANAGER ROUTES ---
# [unchanged]

# --- BOOTSTRAP DB ON FIRST RUN ---
if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
